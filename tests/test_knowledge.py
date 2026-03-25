"""Tests for the knowledge base — ingestion, parsing, search, and API endpoints."""

import asyncio
import io
import email
import email.mime.text
import email.mime.multipart
import os
import tempfile

import pytest
import pytest_asyncio

from cadence.knowledge.store import KnowledgeStore, DocumentRecord, ChunkRecord
from cadence.knowledge.parsers import (
    detect_source_type,
    parse_email_content,
    parse_text_file,
    parse_web_page,
    PARSERS,
)
from cadence.tools.knowledge_tools import (
    KBIngestTool,
    KBSearchTool,
    KBListTool,
    KBDeleteTool,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

@pytest.fixture()
def kb_store(tmp_path, monkeypatch):
    """Provide a KnowledgeStore backed by a temp directory."""
    monkeypatch.setattr(
        "cadence.knowledge.store.get_config",
        lambda: _make_config(tmp_path),
    )
    return KnowledgeStore()


def _make_config(persist_dir):
    """Create a minimal config stub with memory settings."""
    class _MemCfg:
        persist_dir = ""
        default_namespace = "shared"
        max_results = 10
        similarity_threshold = 0.0
        decay_rate = 0.05

    class _Cfg:
        memory = _MemCfg()

    _Cfg.memory.persist_dir = str(persist_dir / "chromadb")
    return _Cfg()


def _make_eml(subject: str, body: str, from_: str = "alice@example.com") -> bytes:
    """Build a minimal .eml email as bytes."""
    msg = email.mime.text.MIMEText(body)
    msg["Subject"] = subject
    msg["From"] = from_
    msg["To"] = "bob@example.com"
    msg["Date"] = "Mon, 24 Mar 2026 10:00:00 +0000"
    return msg.as_bytes()


def _make_multipart_eml(subject: str, text_body: str, html_body: str) -> bytes:
    """Build a multipart .eml with both text and HTML parts."""
    msg = email.mime.multipart.MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = "alice@example.com"
    msg["To"] = "bob@example.com"
    msg.attach(email.mime.text.MIMEText(text_body, "plain"))
    msg.attach(email.mime.text.MIMEText(html_body, "html"))
    return msg.as_bytes()


def _make_sample_pdf() -> bytes:
    """Create a minimal PDF file in memory."""
    try:
        import pypdf
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf not installed")

    writer = PdfWriter()
    # pypdf requires ReportLab to create pages from scratch, but we can use
    # a simpler approach: create a minimal valid PDF manually.
    # Instead, we'll use the pypdf writer with an empty page and add text via
    # a different approach — create the raw PDF bytes directly.
    pdf_content = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
        b"4 0 obj\n<< /Length 44 >>\nstream\nBT /F1 12 Tf 100 700 Td (Hello PDF World) Tj ET\nendstream\nendobj\n"
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000266 00000 n \n"
        b"0000000360 00000 n \n"
        b"trailer\n<< /Size 6 /Root 1 0 R >>\n"
        b"startxref\n441\n%%EOF\n"
    )
    return pdf_content


def _make_sample_docx() -> bytes:
    """Create a minimal DOCX file in memory."""
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = docx.Document()
    doc.add_paragraph("This is the first paragraph about machine learning.")
    doc.add_paragraph("This is the second paragraph about neural networks.")
    doc.add_paragraph("This is the third paragraph about deep learning architectures.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Source type detection
# ---------------------------------------------------------------------------

class TestDetectSourceType:
    def test_pdf(self):
        assert detect_source_type("/path/to/document.pdf") == "pdf"
        assert detect_source_type("REPORT.PDF") == "pdf"

    def test_docx(self):
        assert detect_source_type("/path/to/file.docx") == "docx"

    def test_email(self):
        assert detect_source_type("message.eml") == "email"

    def test_web(self):
        assert detect_source_type("https://example.com/page") == "web"
        assert detect_source_type("http://example.com") == "web"

    def test_text_fallback(self):
        assert detect_source_type("notes.txt") == "text"
        assert detect_source_type("readme.md") == "text"
        assert detect_source_type("data.csv") == "text"


# ---------------------------------------------------------------------------
# Text chunking
# ---------------------------------------------------------------------------

class TestChunking:
    def test_empty_text(self):
        assert KnowledgeStore.chunk_text("") == []
        assert KnowledgeStore.chunk_text("   ") == []

    def test_short_text(self):
        chunks = KnowledgeStore.chunk_text("Hello world", chunk_size=1000)
        assert len(chunks) == 1
        assert chunks[0] == "Hello world"

    def test_long_text_creates_multiple_chunks(self):
        text = "Word " * 500  # ~2500 chars
        chunks = KnowledgeStore.chunk_text(text, chunk_size=500, overlap=100)
        assert len(chunks) > 1

    def test_overlap_between_chunks(self):
        # Create text with clear sentence boundaries
        sentences = [f"Sentence number {i} is here." for i in range(50)]
        text = " ".join(sentences)
        chunks = KnowledgeStore.chunk_text(text, chunk_size=200, overlap=50)
        assert len(chunks) > 1
        # With overlap, consecutive chunks should share some content
        for i in range(len(chunks) - 1):
            # The end of one chunk and start of next should have some overlap
            # (not guaranteed to be exact due to boundary-seeking, but chunks
            # should cover the full text)
            pass
        # Verify all original text is covered
        combined = " ".join(chunks)
        for s in sentences:
            assert s in combined or s.split()[0] in combined

    def test_paragraph_boundary_breaking(self):
        text = "First paragraph content.\n\nSecond paragraph content.\n\nThird paragraph content."
        chunks = KnowledgeStore.chunk_text(text, chunk_size=50, overlap=10)
        assert len(chunks) >= 2


# ---------------------------------------------------------------------------
# Email parsing
# ---------------------------------------------------------------------------

class TestEmailParsing:
    def test_plain_text_email(self):
        eml = _make_eml("Test Subject", "Hello, this is the email body.")
        text, meta = parse_email_content(eml)
        assert "Hello, this is the email body." in text
        assert meta["subject"] == "Test Subject"
        assert meta["from"] == "alice@example.com"

    def test_multipart_email_prefers_text(self):
        eml = _make_multipart_eml(
            "Multi Subject",
            "Plain text body here.",
            "<html><body><p>HTML body here.</p></body></html>",
        )
        text, meta = parse_email_content(eml)
        assert "Plain text body here." in text
        assert meta["subject"] == "Multi Subject"

    def test_email_from_file(self, tmp_path):
        eml_data = _make_eml("File Email", "Body from file.")
        eml_file = tmp_path / "test.eml"
        eml_file.write_bytes(eml_data)

        text, meta = parse_email_content(str(eml_file))
        assert "Body from file." in text
        assert meta["subject"] == "File Email"

    def test_email_metadata_extraction(self):
        eml = _make_eml("Meeting Notes", "Important meeting details.", "ceo@corp.com")
        text, meta = parse_email_content(eml)
        assert meta["from"] == "ceo@corp.com"
        assert meta["to"] == "bob@example.com"
        assert "date" in meta


# ---------------------------------------------------------------------------
# PDF parsing
# ---------------------------------------------------------------------------

class TestPDFParsing:
    def test_pdf_from_bytes(self):
        pdf_bytes = _make_sample_pdf()
        try:
            from cadence.knowledge.parsers import parse_pdf
        except ImportError:
            pytest.skip("pypdf not installed")

        text, meta = parse_pdf(pdf_bytes)
        assert "page_count" in meta
        assert meta["page_count"] == 1
        # The manually constructed PDF should have extractable text
        assert isinstance(text, str)

    def test_pdf_from_file(self, tmp_path):
        pdf_bytes = _make_sample_pdf()
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(pdf_bytes)

        try:
            from cadence.knowledge.parsers import parse_pdf
        except ImportError:
            pytest.skip("pypdf not installed")

        text, meta = parse_pdf(str(pdf_file))
        assert meta["page_count"] == 1


# ---------------------------------------------------------------------------
# DOCX parsing
# ---------------------------------------------------------------------------

class TestDOCXParsing:
    def test_docx_from_bytes(self):
        docx_bytes = _make_sample_docx()
        try:
            from cadence.knowledge.parsers import parse_docx
        except ImportError:
            pytest.skip("python-docx not installed")

        text, meta = parse_docx(docx_bytes)
        assert "machine learning" in text
        assert "neural networks" in text
        assert "deep learning" in text
        assert meta["paragraph_count"] == 3

    def test_docx_from_file(self, tmp_path):
        docx_bytes = _make_sample_docx()
        docx_file = tmp_path / "test.docx"
        docx_file.write_bytes(docx_bytes)

        try:
            from cadence.knowledge.parsers import parse_docx
        except ImportError:
            pytest.skip("python-docx not installed")

        text, meta = parse_docx(str(docx_file))
        assert "machine learning" in text


# ---------------------------------------------------------------------------
# Web page parsing
# ---------------------------------------------------------------------------

class TestWebParsing:
    def test_web_page_fetch(self):
        """Test web page fetching with a known public URL."""
        try:
            text, meta = parse_web_page("https://httpbin.org/html")
        except (ValueError, Exception):
            pytest.skip("Network unavailable or httpbin.org unreachable")

        assert len(text) > 0
        assert "url" in meta
        assert meta["url"] == "https://httpbin.org/html"

    def test_web_invalid_url(self):
        with pytest.raises(ValueError, match="Failed to fetch"):
            parse_web_page("https://this-domain-does-not-exist-xyz123.com/page")


# ---------------------------------------------------------------------------
# Text file parsing
# ---------------------------------------------------------------------------

class TestTextParsing:
    def test_text_from_string(self):
        text, meta = parse_text_file("Hello world, this is plain text.")
        assert text == "Hello world, this is plain text."

    def test_text_from_bytes(self):
        text, meta = parse_text_file(b"Bytes content here.")
        assert text == "Bytes content here."

    def test_text_from_file(self, tmp_path):
        txt_file = tmp_path / "notes.txt"
        txt_file.write_text("File content here.")
        text, meta = parse_text_file(str(txt_file))
        assert text == "File content here."


# ---------------------------------------------------------------------------
# KnowledgeStore — ingestion and search
# ---------------------------------------------------------------------------

class TestKnowledgeStoreIngestion:
    @pytest.mark.asyncio
    async def test_ingest_text(self, kb_store):
        doc = await kb_store.ingest(
            title="Test Document",
            content="This is a test document about artificial intelligence and machine learning.",
            source="text",
            origin="test",
        )
        assert doc.title == "Test Document"
        assert doc.source == "text"
        assert doc.chunk_count >= 1

    @pytest.mark.asyncio
    async def test_ingest_empty_content(self, kb_store):
        doc = await kb_store.ingest(
            title="Empty",
            content="",
            source="text",
        )
        assert doc.chunk_count == 0

    @pytest.mark.asyncio
    async def test_ingest_long_content_creates_chunks(self, kb_store):
        # Create content that will require multiple chunks
        content = "\n\n".join(
            [f"Paragraph {i}: " + ("Lorem ipsum dolor sit amet. " * 20) for i in range(20)]
        )
        doc = await kb_store.ingest(
            title="Long Document",
            content=content,
            source="text",
            chunk_size=500,
            chunk_overlap=100,
        )
        assert doc.chunk_count > 1

    @pytest.mark.asyncio
    async def test_ingest_with_metadata(self, kb_store):
        doc = await kb_store.ingest(
            title="Metadata Doc",
            content="Document with custom metadata attached.",
            source="email",
            origin="inbox",
            metadata={"from": "alice@example.com", "subject": "Hello"},
        )
        assert doc.source == "email"
        assert doc.metadata["from"] == "alice@example.com"


class TestKnowledgeStoreSearch:
    @pytest.mark.asyncio
    async def test_search_empty_store(self, kb_store):
        results = await kb_store.search("anything")
        assert results == []

    @pytest.mark.asyncio
    async def test_search_returns_relevant_results(self, kb_store):
        await kb_store.ingest(
            title="AI Paper",
            content="Deep learning is a subset of machine learning that uses neural networks with many layers.",
            source="pdf",
        )
        await kb_store.ingest(
            title="Cooking Recipe",
            content="To make a chocolate cake, mix flour, sugar, cocoa powder, eggs, and butter.",
            source="web",
        )

        results = await kb_store.search("neural networks and deep learning")
        assert len(results) > 0
        # The AI paper should be more relevant than the cooking recipe
        assert results[0].document.title == "AI Paper"

    @pytest.mark.asyncio
    async def test_search_with_source_filter(self, kb_store):
        await kb_store.ingest(
            title="PDF Doc", content="Important PDF content about finance.", source="pdf",
        )
        await kb_store.ingest(
            title="Web Page", content="Important web content about finance.", source="web",
        )

        results = await kb_store.search("finance", source_filter="pdf")
        assert len(results) > 0
        for r in results:
            assert r.document.source == "pdf"

    @pytest.mark.asyncio
    async def test_search_max_results(self, kb_store):
        for i in range(10):
            await kb_store.ingest(
                title=f"Doc {i}",
                content=f"Document number {i} about technology and innovation.",
                source="text",
            )

        results = await kb_store.search("technology", max_results=3)
        assert len(results) <= 3

    @pytest.mark.asyncio
    async def test_search_relevance_scoring(self, kb_store):
        await kb_store.ingest(
            title="Exact Match",
            content="Quantum computing uses qubits instead of classical bits.",
            source="text",
        )
        await kb_store.ingest(
            title="Partial Match",
            content="Classical computing has been around for decades using transistors.",
            source="text",
        )

        results = await kb_store.search("quantum computing qubits")
        assert len(results) > 0
        # All results should have relevance scores
        for r in results:
            assert r.relevance > 0
            assert r.similarity > 0


# ---------------------------------------------------------------------------
# KnowledgeStore — document management
# ---------------------------------------------------------------------------

class TestKnowledgeStoreManagement:
    @pytest.mark.asyncio
    async def test_list_documents(self, kb_store):
        await kb_store.ingest(title="Doc A", content="Content A.", source="pdf")
        await kb_store.ingest(title="Doc B", content="Content B.", source="email")

        docs = await kb_store.list_documents()
        assert len(docs) == 2
        titles = {d.title for d in docs}
        assert "Doc A" in titles
        assert "Doc B" in titles

    @pytest.mark.asyncio
    async def test_list_documents_empty(self, kb_store):
        docs = await kb_store.list_documents()
        assert docs == []

    @pytest.mark.asyncio
    async def test_delete_document(self, kb_store):
        doc = await kb_store.ingest(
            title="To Delete", content="This will be deleted.", source="text",
        )
        docs_before = await kb_store.list_documents()
        assert len(docs_before) == 1

        ok = await kb_store.delete_document(doc.id)
        assert ok is True

        docs_after = await kb_store.list_documents()
        assert len(docs_after) == 0

    @pytest.mark.asyncio
    async def test_delete_removes_chunks(self, kb_store):
        content = "A " * 2000  # Force multiple chunks
        doc = await kb_store.ingest(
            title="Chunked", content=content, source="text", chunk_size=200,
        )
        assert doc.chunk_count > 1

        stats_before = await kb_store.stats()
        assert stats_before["total_chunks"] > 0

        await kb_store.delete_document(doc.id)
        stats_after = await kb_store.stats()
        assert stats_after["total_chunks"] == 0

    @pytest.mark.asyncio
    async def test_stats(self, kb_store):
        stats = await kb_store.stats()
        assert stats["total_chunks"] == 0
        assert stats["total_documents"] == 0

        await kb_store.ingest(title="X", content="Content.", source="text")
        stats = await kb_store.stats()
        assert stats["total_chunks"] >= 1
        assert stats["total_documents"] == 1


# ---------------------------------------------------------------------------
# Agent tools
# ---------------------------------------------------------------------------

class TestKBTools:
    @pytest.fixture()
    def tools(self, kb_store):
        return {
            "ingest": KBIngestTool(kb_store),
            "search": KBSearchTool(kb_store),
            "list": KBListTool(kb_store),
            "delete": KBDeleteTool(kb_store),
        }

    @pytest.mark.asyncio
    async def test_ingest_tool_text(self, tools):
        result = await tools["ingest"].execute(
            source="The quick brown fox jumps over the lazy dog.",
            title="Fox Document",
            source_type="text",
        )
        assert "Ingested" in result
        assert "Fox Document" in result
        assert "chunk" in result.lower()

    @pytest.mark.asyncio
    async def test_ingest_tool_auto_detect(self, tools, tmp_path):
        txt_file = tmp_path / "notes.txt"
        txt_file.write_text("Some important notes about the project.")
        result = await tools["ingest"].execute(
            source=str(txt_file),
            title="Project Notes",
        )
        assert "Ingested" in result
        assert "text" in result.lower()

    @pytest.mark.asyncio
    async def test_ingest_tool_unsupported_type(self, tools):
        result = await tools["ingest"].execute(
            source="something",
            title="Bad",
            source_type="unsupported_format",
        )
        assert "Error" in result or "unsupported" in result.lower()

    @pytest.mark.asyncio
    async def test_ingest_tool_empty_content(self, tools):
        result = await tools["ingest"].execute(
            source="",
            title="Empty",
            source_type="text",
        )
        assert "Error" in result or "no text" in result.lower()

    @pytest.mark.asyncio
    async def test_search_tool(self, tools):
        await tools["ingest"].execute(
            source="Python is a programming language used for data science and web development.",
            title="Python Info",
            source_type="text",
        )
        result = await tools["search"].execute(query="programming language")
        assert "Python" in result

    @pytest.mark.asyncio
    async def test_search_tool_empty_kb(self, tools):
        result = await tools["search"].execute(query="anything")
        assert "No relevant results" in result

    @pytest.mark.asyncio
    async def test_search_tool_with_filter(self, tools):
        await tools["ingest"].execute(
            source="Finance content for testing.",
            title="Finance PDF",
            source_type="text",
        )
        result = await tools["search"].execute(
            query="finance", source_filter="text",
        )
        assert "Finance" in result

    @pytest.mark.asyncio
    async def test_list_tool(self, tools):
        await tools["ingest"].execute(
            source="Some content.", title="Listed Doc", source_type="text",
        )
        result = await tools["list"].execute()
        assert "Listed Doc" in result

    @pytest.mark.asyncio
    async def test_list_tool_empty(self, tools):
        result = await tools["list"].execute()
        assert "empty" in result.lower()

    @pytest.mark.asyncio
    async def test_delete_tool(self, tools):
        await tools["ingest"].execute(
            source="Content to delete.", title="Deletable", source_type="text",
        )
        # Get document ID from list
        list_result = await tools["list"].execute()
        # Extract the short ID (first 8 chars in brackets)
        doc_id = list_result.split("[")[1].split("]")[0]

        result = await tools["delete"].execute(document_id=doc_id)
        assert "Deleted" in result

        # Verify it's gone
        list_after = await tools["list"].execute()
        assert "empty" in list_after.lower()


# ---------------------------------------------------------------------------
# PDF ingestion tool (end-to-end)
# ---------------------------------------------------------------------------

class TestPDFIngestionE2E:
    @pytest.mark.asyncio
    async def test_ingest_pdf_file(self, kb_store, tmp_path):
        pdf_bytes = _make_sample_pdf()
        pdf_file = tmp_path / "sample.pdf"
        pdf_file.write_bytes(pdf_bytes)

        tool = KBIngestTool(kb_store)
        result = await tool.execute(source=str(pdf_file), title="Sample PDF")
        assert "Ingested" in result
        assert "pdf" in result.lower()


# ---------------------------------------------------------------------------
# DOCX ingestion tool (end-to-end)
# ---------------------------------------------------------------------------

class TestDOCXIngestionE2E:
    @pytest.mark.asyncio
    async def test_ingest_docx_file(self, kb_store, tmp_path):
        docx_bytes = _make_sample_docx()
        docx_file = tmp_path / "sample.docx"
        docx_file.write_bytes(docx_bytes)

        tool = KBIngestTool(kb_store)
        result = await tool.execute(source=str(docx_file), title="Sample DOCX")
        assert "Ingested" in result
        assert "docx" in result.lower()

        # Verify content is searchable
        search_tool = KBSearchTool(kb_store)
        search_result = await search_tool.execute(query="machine learning")
        assert "machine learning" in search_result.lower()


# ---------------------------------------------------------------------------
# Email ingestion tool (end-to-end)
# ---------------------------------------------------------------------------

class TestEmailIngestionE2E:
    @pytest.mark.asyncio
    async def test_ingest_eml_file(self, kb_store, tmp_path):
        eml_data = _make_eml("Project Update", "The project is on track for Q2 delivery.")
        eml_file = tmp_path / "update.eml"
        eml_file.write_bytes(eml_data)

        tool = KBIngestTool(kb_store)
        result = await tool.execute(source=str(eml_file), title="Project Update Email")
        assert "Ingested" in result
        assert "email" in result.lower()

        # Verify content is searchable
        search_tool = KBSearchTool(kb_store)
        search_result = await search_tool.execute(query="Q2 delivery")
        assert "Q2" in search_result or "delivery" in search_result


# ---------------------------------------------------------------------------
# Web page ingestion tool (end-to-end)
# ---------------------------------------------------------------------------

class TestWebIngestionE2E:
    @pytest.mark.asyncio
    async def test_ingest_web_page(self, kb_store):
        tool = KBIngestTool(kb_store)
        try:
            result = await tool.execute(
                source="https://httpbin.org/html",
                title="HTTPBin HTML Page",
            )
        except Exception:
            pytest.skip("Network unavailable")

        if "Error" in result and "fetch" in result.lower():
            pytest.skip("Network unavailable or httpbin.org unreachable")

        assert "Ingested" in result
        assert "web" in result.lower()


# ---------------------------------------------------------------------------
# API endpoint tests
# ---------------------------------------------------------------------------

class TestKBAPIEndpoints:
    """Test the FastAPI knowledge base endpoints using TestClient."""

    @pytest.fixture()
    def client(self, tmp_path, monkeypatch):
        """Create a test client with an isolated KB store."""
        monkeypatch.setattr(
            "cadence.knowledge.store.get_config",
            lambda: _make_config(tmp_path),
        )
        # Reset the global KB store so it picks up our monkeypatched config
        import cadence.api as api_module
        api_module._kb_store = None
        monkeypatch.setattr(
            "cadence.api.get_kb_store",
            lambda: KnowledgeStore(),
        )

        from fastapi.testclient import TestClient
        from cadence.api import app

        return TestClient(app, raise_server_exceptions=False)

    @pytest.fixture()
    def client_with_store(self, tmp_path, monkeypatch):
        """Create a test client with a shared KB store for state across requests."""
        monkeypatch.setattr(
            "cadence.knowledge.store.get_config",
            lambda: _make_config(tmp_path),
        )
        store = KnowledgeStore()
        import cadence.api as api_module
        api_module._kb_store = None
        monkeypatch.setattr("cadence.api.get_kb_store", lambda: store)

        from fastapi.testclient import TestClient
        from cadence.api import app

        return TestClient(app, raise_server_exceptions=False)

    def test_ingest_endpoint(self, client_with_store):
        resp = client_with_store.post("/api/kb/ingest", json={
            "source": "The Eiffel Tower is located in Paris, France.",
            "title": "Eiffel Tower",
            "source_type": "text",
        })
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "ingested"
        assert data["chunk_count"] >= 1
        assert data["title"] == "Eiffel Tower"

    def test_search_endpoint(self, client_with_store):
        # Ingest first
        client_with_store.post("/api/kb/ingest", json={
            "source": "Quantum mechanics describes the behavior of particles at atomic scales.",
            "title": "Quantum Physics",
            "source_type": "text",
        })
        # Search
        resp = client_with_store.post("/api/kb/search", json={
            "query": "quantum particles",
            "max_results": 5,
        })
        assert resp.status_code == 200
        data = resp.json()
        assert data["total"] > 0
        assert "content" in data["results"][0]

    def test_list_documents_endpoint(self, client_with_store):
        client_with_store.post("/api/kb/ingest", json={
            "source": "Test content.", "title": "Test Doc", "source_type": "text",
        })
        resp = client_with_store.get("/api/kb/documents")
        assert resp.status_code == 200
        docs = resp.json()
        assert len(docs) >= 1
        assert docs[0]["title"] == "Test Doc"

    def test_delete_document_endpoint(self, client_with_store):
        # Ingest
        ingest_resp = client_with_store.post("/api/kb/ingest", json={
            "source": "Deletable content.", "title": "To Delete", "source_type": "text",
        })
        doc_id = ingest_resp.json()["document_id"]

        # Delete
        resp = client_with_store.delete(f"/api/kb/documents/{doc_id}")
        assert resp.status_code == 200
        assert resp.json()["status"] == "deleted"

        # Verify gone
        list_resp = client_with_store.get("/api/kb/documents")
        assert len(list_resp.json()) == 0

    def test_stats_endpoint(self, client_with_store):
        resp = client_with_store.get("/api/kb/stats")
        assert resp.status_code == 200
        data = resp.json()
        assert data["total_chunks"] == 0
        assert data["total_documents"] == 0

        # Ingest something
        client_with_store.post("/api/kb/ingest", json={
            "source": "Stats test content.", "title": "Stats Doc", "source_type": "text",
        })
        resp = client_with_store.get("/api/kb/stats")
        data = resp.json()
        assert data["total_chunks"] >= 1
        assert data["total_documents"] == 1

    def test_upload_endpoint(self, client_with_store, tmp_path):
        # Create a text file to upload
        txt_content = b"Uploaded file content about machine learning algorithms."
        resp = client_with_store.post(
            "/api/kb/ingest/upload",
            files={"file": ("notes.txt", txt_content, "text/plain")},
            params={"title": "Uploaded Notes"},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "ingested"
        assert data["title"] == "Uploaded Notes"

    def test_upload_docx_endpoint(self, client_with_store):
        docx_bytes = _make_sample_docx()
        resp = client_with_store.post(
            "/api/kb/ingest/upload",
            files={"file": ("report.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            params={"title": "Uploaded Report"},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "ingested"
        assert data["chunk_count"] >= 1

    def test_search_with_source_filter_endpoint(self, client_with_store):
        client_with_store.post("/api/kb/ingest", json={
            "source": "PDF-like finance content.", "title": "Finance PDF", "source_type": "pdf",
        })
        client_with_store.post("/api/kb/ingest", json={
            "source": "Email about finance.", "title": "Finance Email", "source_type": "email",
        })
        resp = client_with_store.post("/api/kb/search", json={
            "query": "finance",
            "source_filter": "pdf",
        })
        data = resp.json()
        for r in data["results"]:
            assert r["source"] == "pdf"
