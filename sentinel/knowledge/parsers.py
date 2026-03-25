"""Document parsers — extract text from PDFs, DOCX, emails, and web pages."""

from __future__ import annotations

import email
import email.policy
import io
import urllib.request
import urllib.error
from html.parser import HTMLParser
from pathlib import Path
from typing import Any


class _TextExtractor(HTMLParser):
    """Minimal HTML-to-text extractor."""

    def __init__(self):
        super().__init__()
        self._chunks: list[str] = []
        self._skip = False

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style", "noscript"):
            self._skip = True

    def handle_endtag(self, tag):
        if tag in ("script", "style", "noscript"):
            self._skip = False

    def handle_data(self, data):
        if not self._skip:
            text = data.strip()
            if text:
                self._chunks.append(text)

    def get_text(self) -> str:
        return "\n".join(self._chunks)


def parse_pdf(source: str | bytes) -> tuple[str, dict[str, Any]]:
    """Extract text from a PDF file path or raw bytes.

    Returns (text, metadata).
    """
    try:
        import pypdf
    except ImportError:
        raise ImportError("pypdf is required for PDF ingestion. Install with: pip install pypdf")

    if isinstance(source, bytes):
        reader = pypdf.PdfReader(io.BytesIO(source))
    else:
        reader = pypdf.PdfReader(source)

    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)

    metadata: dict[str, Any] = {"page_count": len(reader.pages)}
    if reader.metadata:
        if reader.metadata.title:
            metadata["pdf_title"] = reader.metadata.title
        if reader.metadata.author:
            metadata["pdf_author"] = reader.metadata.author

    return "\n\n".join(pages), metadata


def parse_docx(source: str | bytes) -> tuple[str, dict[str, Any]]:
    """Extract text from a DOCX file path or raw bytes.

    Returns (text, metadata).
    """
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX ingestion. Install with: pip install python-docx"
        )

    if isinstance(source, bytes):
        document = docx.Document(io.BytesIO(source))
    else:
        document = docx.Document(source)

    paragraphs = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    metadata: dict[str, Any] = {"paragraph_count": len(paragraphs)}
    core = document.core_properties
    if core.title:
        metadata["docx_title"] = core.title
    if core.author:
        metadata["docx_author"] = core.author

    return "\n\n".join(paragraphs), metadata


def parse_email_content(source: str | bytes) -> tuple[str, dict[str, Any]]:
    """Parse an email (.eml) from a file path, raw bytes, or string content.

    Returns (text_body, metadata).
    """
    if isinstance(source, (str, Path)) and Path(source).exists():
        with open(source, "rb") as f:
            raw = f.read()
    elif isinstance(source, bytes):
        raw = source
    else:
        raw = source.encode("utf-8") if isinstance(source, str) else source

    msg = email.message_from_bytes(raw, policy=email.policy.default)

    metadata: dict[str, Any] = {}
    for header in ("subject", "from", "to", "date"):
        val = msg.get(header)
        if val:
            metadata[header] = str(val)

    # Extract text body
    body_parts = []
    if msg.is_multipart():
        for part in msg.walk():
            ct = part.get_content_type()
            if ct == "text/plain":
                payload = part.get_content()
                if payload:
                    body_parts.append(str(payload))
            elif ct == "text/html" and not body_parts:
                payload = part.get_content()
                if payload:
                    extractor = _TextExtractor()
                    extractor.feed(str(payload))
                    body_parts.append(extractor.get_text())
    else:
        ct = msg.get_content_type()
        payload = msg.get_content()
        if payload:
            if ct == "text/html":
                extractor = _TextExtractor()
                extractor.feed(str(payload))
                body_parts.append(extractor.get_text())
            else:
                body_parts.append(str(payload))

    return "\n\n".join(body_parts), metadata


def parse_web_page(url: str, max_chars: int = 100000) -> tuple[str, dict[str, Any]]:
    """Fetch and extract text from a web page URL.

    Returns (text, metadata).
    """
    req = urllib.request.Request(url, headers={"User-Agent": "Sentinel/0.1"})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            content_type = resp.headers.get("Content-Type", "")
            raw = resp.read().decode(errors="replace")
    except urllib.error.URLError as e:
        raise ValueError(f"Failed to fetch {url}: {e}")

    metadata: dict[str, Any] = {"url": url, "content_type": content_type}

    if "html" in content_type.lower():
        extractor = _TextExtractor()
        extractor.feed(raw)
        text = extractor.get_text()
    else:
        text = raw

    if len(text) > max_chars:
        text = text[:max_chars]

    return text, metadata


def parse_text_file(source: str | bytes) -> tuple[str, dict[str, Any]]:
    """Read a plain text file.

    Returns (text, metadata).
    """
    if isinstance(source, bytes):
        text = source.decode(errors="replace")
    elif Path(source).exists():
        text = Path(source).read_text(errors="replace")
    else:
        text = source

    return text, {"type": "text"}


# Dispatcher mapping source types to parsers
PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "email": parse_email_content,
    "web": parse_web_page,
    "text": parse_text_file,
}


def detect_source_type(path_or_url: str) -> str:
    """Auto-detect source type from file extension or URL."""
    lower = path_or_url.lower()
    if lower.startswith("http://") or lower.startswith("https://"):
        return "web"
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith(".eml") or lower.endswith(".msg"):
        return "email"
    return "text"
